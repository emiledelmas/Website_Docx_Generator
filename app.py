from flask import Flask, request, render_template, redirect, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import os
# Import Wikipedia library to get biography of writers
import wikipedia
# Change the langage below to "en" if you are not french
wikipedia.set_lang("fr")
app = Flask(__name__)


@app.route('/')
def my_form():
    return render_template('index.html') # Return here your Html Page (with the form)


@app.route('/', methods=['POST']) # When the user click to the "submit" button
def my_form_post():
# First we put into variables all the data that the user entered
    titre = request.form['titre']
    author = request.form['auteur']
    number = request.form['number']
    problematique = request.form['problematique']
    nbrePartie = request.form['nbrePartie']
    I = request.form['I']
    I1 = request.form['I1']
    I2 = request.form['I2']
    I3 = request.form['I3']
    II = request.form['II']
    II1 = request.form['II1']
    II2 = request.form['II2']
    II3 = request.form['II3']
    III = request.form['III']
    III1 = request.form['III1']
    III2 = request.form['III2']
    III3 = request.form['III3']

# Then an if to know if there is 3 column or 2 because it changes the docx file
    if nbrePartie == '3':
        document = Document('templates-docx/Template3.docx')

# We created a function that deleted paragraph in tables already created
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

        try:
            search = wikipedia.search(author)
            author = search[0]
            getsummary = document.tables[6]
            delete_paragraph(getsummary.rows[0].cells[0].paragraphs[-1])
            getsummary.rows[0].cells[0].add_paragraph(wikipedia.summary(author, sentences=1))
        except wikipedia.exceptions.DisambiguationError as e:
            print("a")

        getTitle = document.tables[0]
        delete_paragraph(getTitle.rows[0].cells[0].paragraphs[-1])
        getTitle.rows[0].cells[0].add_paragraph(titre, style="GrandTitre")
        getAuthor = document.tables[4]
        delete_paragraph(getAuthor.rows[0].cells[0].paragraphs[-1])
        getAuthor.rows[0].cells[0].add_paragraph("De "+author, style="auteur")
        getNumber = document.tables[1]
        delete_paragraph(getNumber.rows[0].cells[0].paragraphs[-1])
        getNumber.rows[0].cells[0].add_paragraph(number, style="number")
        getPb = document.tables[2]
        delete_paragraph(getPb.rows[0].cells[0].paragraphs[-1])
        getPb.rows[0].cells[0].add_paragraph(problematique)
        getPlan = document.tables[5]
        delete_paragraph(getPlan.rows[0].cells[0].paragraphs[-1])
        getPlan.rows[0].cells[0].add_paragraph("I. "+I, style="PlanG")
        getPlan.rows[0].cells[0].add_paragraph("1. " + I1, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("2. " + I2, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("3. " + I3, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("II. " + II, style="PlanG")
        getPlan.rows[0].cells[0].add_paragraph("1. " + II1, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("2. " + II2, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("3. " + II3, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("III. " + III, style="PlanG")
        getPlan.rows[0].cells[0].add_paragraph("1. " + III1, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("2. " + III2, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("3. " + III3, style="planP")
        plan = (
            (str('I. ')+I, I1, I2, I3),
            (str('II. ')+II, II1, II2, II3),
            (str('III. ')+III, III1, III2, III3),
        )
        table = document.tables[3]

        i = 0
        while i < 4:
            y = 0
            while y < 4:
                delete_paragraph(table.rows[i].cells[y].paragraphs[-1])
                y += 1
            i += 1

        hdr_cells = table.rows[0].cells
        parties = hdr_cells[0].add_paragraph('Parties')
        parties.alignment = WD_ALIGN_PARAGRAPH.CENTER
        one = hdr_cells[1].add_paragraph('1')
        one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        two = hdr_cells[2].add_paragraph('2')
        two.alignment = WD_ALIGN_PARAGRAPH.CENTER
        three = hdr_cells[3].add_paragraph('3')
        three.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i = 0
        for partie, sspartie1, sspartie2, sspartie3 in plan:
            i += 1;
            hdr_cells = table.rows[i].cells

            new_partie = hdr_cells[0].add_paragraph(partie, style="Titre1")
            new_partie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            new_sspartie1 = hdr_cells[1].add_paragraph(sspartie1, style="titressparties")
            new_sspartie1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            new_sspartie2 = hdr_cells[2].add_paragraph(sspartie2, style="titressparties")
            new_sspartie2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            new_sspartie3 = hdr_cells[3].add_paragraph(sspartie3, style="titressparties")
            new_sspartie3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        root = "Render/"+str(number)+'.'+str(author)+'-'+str(titre)+'.docx'
        name = str(number)+'. '+str(author)+' - '+str(titre)+'.docx'
        document.save(root)
        return send_file(root,name,as_attachment=True,
                                    attachment_filename=os.path.basename(name))
    else:
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None


        document = Document('templates-docx/Template2.docx')
        print(document.tables)
        getTitle = document.tables[0]
        delete_paragraph(getTitle.rows[0].cells[0].paragraphs[-1])
        getTitle.rows[0].cells[0].add_paragraph(titre, style="GrandTitre")
        getAuthor = document.tables[4]
        delete_paragraph(getAuthor.rows[0].cells[0].paragraphs[-1])
        getAuthor.rows[0].cells[0].add_paragraph("De "+author, style="auteur")
        getNumber = document.tables[1]
        delete_paragraph(getNumber.rows[0].cells[0].paragraphs[-1])
        getNumber.rows[0].cells[0].add_paragraph(number, style="number")
        getPb = document.tables[2]
        delete_paragraph(getPb.rows[0].cells[0].paragraphs[-1])
        getPb.rows[0].cells[0].add_paragraph(problematique)
        getPlan = document.tables[5]
        delete_paragraph(getPlan.rows[0].cells[0].paragraphs[-1])
        getPlan.rows[0].cells[0].add_paragraph("I. "+I, style="PlanG")
        getPlan.rows[0].cells[0].add_paragraph("1. " + I1, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("2. " + I2, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("3. " + I3, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("II. " + II, style="PlanG")
        getPlan.rows[0].cells[0].add_paragraph("1. " + II1, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("2. " + II2, style="planP")
        getPlan.rows[0].cells[0].add_paragraph("3. " + II3, style="planP")
        plan = (
            (str('I. ')+I, I1, I2, I3),
            (str('II. ')+II, II1, II2, II3),
        )
        table = document.tables[3]

        i = 0
        while i < 3:
            y = 0
            while y < 4:
                delete_paragraph(table.rows[i].cells[y].paragraphs[-1])
                y += 1
            i += 1

        hdr_cells = table.rows[0].cells
        parties = hdr_cells[0].add_paragraph('Parties')
        parties.alignment = WD_ALIGN_PARAGRAPH.CENTER
        one = hdr_cells[1].add_paragraph('1')
        one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        two = hdr_cells[2].add_paragraph('2')
        two.alignment = WD_ALIGN_PARAGRAPH.CENTER
        three = hdr_cells[3].add_paragraph('3')
        three.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i = 0
        for partie, sspartie1, sspartie2, sspartie3 in plan:
            i += 1;
            hdr_cells = table.rows[i].cells

            new_partie = hdr_cells[0].add_paragraph(partie, style="Titre1")
            new_partie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            new_sspartie1 = hdr_cells[1].add_paragraph(sspartie1, style="titressparties")
            new_sspartie1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            new_sspartie2 = hdr_cells[2].add_paragraph(sspartie2, style="titressparties")
            new_sspartie2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            new_sspartie3 = hdr_cells[3].add_paragraph(sspartie3, style="titressparties")
            new_sspartie3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        root = "Render/"+str(number)+'.'+str(author)+'-'+str(titre)+'.docx'
        name = str(number)+'. '+str(author)+' - '+str(titre)+'.docx'
        document.save(root)
        return send_file(root,name,as_attachment=True,
                                    attachment_filename=os.path.basename(name))